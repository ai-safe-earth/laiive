"""Convert voice, image, URL, and text inputs into an EventDraft.

The one extraction prompt for every modality (04-plan: the duplicated
prompts/functions merged here). Everything funnels into `EventDraft`.
"""

import base64
import io
import json
from datetime import date

import httpx
from laiive_shared import EventDraft, transcribe
from laiive_shared.drafts import entries_from_json, entry_to_draft, strip_fences
from loguru import logger
from openai import OpenAI

from config import settings

_client = OpenAI(api_key=settings.openai_api_key)

EXTRACTION_PROMPT_VERSION = "v4"

EXTRACTION_PROMPT = """Extract live music events from this text. Today is {today}.

The text may describe ONE event or MANY — a spreadsheet of gigs, a festival
line-up, a promoter listing their season. Return every event you find.

Return JSON of the form {{"events": [ ... ]}}, where each entry carries ONLY
the fields you can actually identify:
{{
  "name": string,            // event title if stated (do NOT invent one)
  "artists": [string, ...],  // performing artists/bands/DJs
  "start_at": "YYYY-MM-DDTHH:MM:SS",  // resolve relative dates using today's date
  "start_at_claim": string,  // the date EXACTLY as the text words it, verbatim, when it names a weekday or writes the date out ("Wednesday 15 September", "sábado 3 de mayo"). Copy it; do not translate, reformat or resolve it. Omit when the text words no date.
  "venue": string,
  "address": string,         // street address if stated
  "city": string,
  "venue_type": "club" | "bar" | "concert_hall" | "arena" | "festival_site" | "open_air" | "other",  // only when the text says so — omit rather than guessing "other"
  "price_min": number,       // 0 for free events
  "price_max": number,       // only when a range is stated
  "price_currency": string,  // ISO code (EUR, USD...) only when stated or implied by symbol
  "description": string,     // short description in the text's own words
  "genre": string,           // lowercase-hyphenated slug, e.g. "indie-rock"
  "ticket_url": string
}}

Rules:
- One entry per event; a single event is a list of one.
- Keep the events in the order they appear in the text, and keep that order
  stable — the same text must always produce the same events in the same
  order, because the promoter refers to them by position ("the third one").
- A detail stated once for the whole set ("all of these are at Sala Clamores")
  belongs on every event it covers.
- Omit any field that is not present. NEVER invent data.
- Numbers for prices — no currency symbols.
- start_at_claim is a QUOTE, not a conclusion. If the text says "Wednesday 15
  September" and the 15th is a Tuesday, still copy "Wednesday 15 September"
  and still resolve start_at to the 15th. Do NOT quietly fix either one: the
  disagreement is the point, and a human is asked about it downstream.
- JSON only. No explanation, no markdown fences.

Text to extract from:
{text}"""


REFINE_PROMPT = """A promoter is completing the details of ONE live music event. Today is {today}.

Current draft (JSON):
{draft}

The promoter's latest message:
{text}

If the message adds or corrects details of THIS event, merge them into the
draft. If it clearly talks about a different event, or carries no event
details at all (a greeting, a confirmation), return the draft unchanged.

Return the full updated draft as a single JSON object with the same fields as
the input, plus any of: name, artists (list), start_at ("YYYY-MM-DDTHH:MM:SS",
resolve relative dates from today), start_at_claim (the date verbatim as the
promoter worded it, when this message states one in words), venue, address,
city, venue_type ("club" | "bar" | "concert_hall" | "arena" |
"festival_site" | "open_air" | "other", only when stated), price_min (number, 0 for free), price_max (number),
price_currency (ISO code, only when stated or implied by symbol), description,
genre (lowercase-hyphenated slug), ticket_url.

Rules:
- NEVER invent data; never drop a field the message did not change.
- Numbers for prices — no currency symbols.
- When this message restates the date, replace start_at_claim with its wording
  and re-resolve start_at from it. When it does not mention the date, leave
  both exactly as they are — a stale claim would be checked against a date it
  was never written beside.
- JSON only. No explanation, no markdown fences."""


def refine_draft(draft: EventDraft, text: str) -> EventDraft:
    """Merge the promoter's latest message into one draft — the walk's turn.

    Mid-walk there is no full re-extraction: only the event under the cursor
    is refined, so a turn costs one draft's worth of tokens however long the
    set is. Anything unparseable falls back to the draft unchanged — a bad
    model reply must never eat the promoter's data.
    """
    response = _client.chat.completions.create(
        model=settings.conversation_model,
        messages=[
            {
                "role": "user",
                "content": REFINE_PROMPT.format(
                    today=date.today().isoformat(),
                    draft=draft.model_dump_json(exclude_none=True),
                    text=text,
                ),
            }
        ],
        temperature=0.0,
    )
    content = strip_fences(response.choices[0].message.content.strip())
    try:
        data = json.loads(content)
    except json.JSONDecodeError:
        logger.warning(f"Failed to parse refine response: {content[:200]}")
        return draft
    if not isinstance(data, dict):
        return draft
    return entry_to_draft(data) or draft


def extract_drafts_from_text(text: str) -> list[EventDraft]:
    """LLM extraction of every event described in free text.

    One extraction path whether the promoter sent one gig or fifty: the whole
    conversation is re-read every turn, so a later sentence ("all of them are
    at Sala Clamores", "the third one starts at 21:00") lands on the right
    drafts without any merge rules on either side of the wire.
    """
    response = _client.chat.completions.create(
        model=settings.conversation_model,
        messages=[
            {
                "role": "user",
                "content": EXTRACTION_PROMPT.format(
                    today=date.today().isoformat(), text=text
                ),
            }
        ],
        temperature=0.0,
    )
    content = strip_fences(response.choices[0].message.content.strip())
    try:
        data = json.loads(content)
    except json.JSONDecodeError:
        logger.warning(f"Failed to parse extraction response: {content[:200]}")
        return []

    drafts = [
        draft
        for entry in entries_from_json(data)
        if (draft := entry_to_draft(entry)) is not None
    ]
    if len(drafts) > 1:
        logger.info(f"Extraction recognized {len(drafts)} events")
    return drafts


def audio_to_text(audio_bytes: bytes, filename: str = "audio.webm") -> str:
    """Transcribe audio using OpenAI Whisper (shared with the retriever)."""
    return transcribe(_client, audio_bytes, filename, settings.whisper_model)


def image_to_text(image_bytes: bytes, mime_type: str = "image/png") -> str:
    """Describe event information found in an image (poster, flyer, ...)."""
    b64 = base64.b64encode(image_bytes).decode()
    response = _client.chat.completions.create(
        model=settings.conversation_model,
        messages=[
            {
                "role": "user",
                "content": [
                    {
                        "type": "text",
                        "text": (
                            "Extract all live music event information from this image. "
                            "Include: event name, artist names, date and time, venue, "
                            "address, city, price, description, genre, and ticket URL "
                            "if visible. Return the information as plain text."
                        ),
                    },
                    {
                        "type": "image_url",
                        "image_url": {"url": f"data:{mime_type};base64,{b64}"},
                    },
                ],
            }
        ],
        temperature=0.0,
    )
    return response.choices[0].message.content.strip()


class UnreadableDocument(ValueError):
    """The document carries no text we can extract."""


def document_to_text(doc_bytes: bytes, filename: str) -> str:
    """Extract text from a document (PDF, DOCX, or plain text).

    PDFs are read through their text layer. A scanned flyer has none, and
    OCR-ing it would mean rendering pages to images (pymupdf, ~40 MB of
    dependency) — the vision path already handles flyers well, so we ask for
    the image instead of carrying the renderer.
    """
    name = filename.lower()

    if name.endswith((".txt", ".csv", ".md", ".json")):
        return doc_bytes.decode("utf-8", errors="replace")

    if name.endswith(".pdf"):
        from pypdf import PdfReader

        try:
            reader = PdfReader(io.BytesIO(doc_bytes))
            text = "\n".join(page.extract_text() or "" for page in reader.pages)
        except Exception as e:
            # A file the user picked badly is their problem to fix, not a 5xx.
            logger.warning(f"Could not parse PDF {filename}: {e}")
            raise UnreadableDocument(
                "Could not read that PDF. Send the flyer as an image instead."
            ) from e
        if not text.strip():
            raise UnreadableDocument(
                "This PDF has no text layer (it looks scanned). "
                "Send the flyer as an image instead."
            )
        return text.strip()

    if name.endswith(".docx"):
        from docx import Document

        try:
            document = Document(io.BytesIO(doc_bytes))
        except Exception as e:
            logger.warning(f"Could not parse DOCX {filename}: {e}")
            raise UnreadableDocument("Could not read that document.") from e
        return "\n".join(p.text for p in document.paragraphs).strip()

    raise UnreadableDocument(
        f"Unsupported document type: {filename}. Send a PDF, DOCX, TXT, or an image."
    )


URL_MAX_CHARS = 8000


def url_to_text(url: str) -> str:
    """Fetch a page and return its content as text, truncated to a sane size."""
    try:
        with httpx.Client(timeout=15.0, follow_redirects=True) as http:
            resp = http.get(
                url,
                headers={"User-Agent": "Mozilla/5.0 (compatible; laiive-bot/1.0)"},
            )
            resp.raise_for_status()
            page_text = resp.text
    except Exception as e:
        logger.error(f"Failed to fetch URL {url}: {e}")
        raise ValueError(f"Could not fetch URL: {e}") from e

    return page_text[:URL_MAX_CHARS]
